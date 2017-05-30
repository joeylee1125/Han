# -*- coding: UTF-8 -*-
import re
import sys
import os
import time
import csv
import argparse
import codecs

from shutil import copytree
from shutil import copyfile
from docx import Document

import Spider
import CourtList
import FileUtils
import DocAnalyser



def get_case_info(wenshu):
    wenshu.getCaseList(wenshu.total_items)
    

            
            
def validate_case(path, court):
    csv_file = path + court + '.csv'
    csv_file_valid = path + court + '_valid.csv'
    case_matrix = FileUtils.read_csv(csv_file)
    case_matrix['valid'] = ['N'] * len(case_matrix['name'])
    
    analyser = DocAnalyser.DocAnalyser()
    #for i in range(len(case_matrix['name'])):
    for i, case_name in enumerate(case_matrix['name']):
        doc_name = path + court + '\\' + case_name + case_matrix['date'][i] + '.docx'
        FileUtils.valid_doc(file_name)
    
        if analyser.bgrt:
            case_matrix['valid'][i] = 'Y'
        else:
            print("%s is invalid" % case_name)
    FileUtils.dump2csv(case_matrix, csv_file_valid)


def download_case_in_zip(search_criteria, court, path):
    folder = path + court
    csv_file = path + court + '.csv'
    print('Read case list from %s' % csv_file)
    # Read csv file and get case list.
    case_matrix = FileUtils.read_csv(csv_file)
    wenshu = Spider.WenShu()
    wenshu.setSearchCriteria(search_criteria)
    
    row_count = len(case_matrix['name'])

    FileUtils.validate_path(folder)
    if row_count % 20 is 0:
        total_page = row_count // 20
        last_page = 20
    else:
        total_page = (row_count // 20) + 1  
        last_page = row_count % 20
    for page in range(total_page):
        
        file_name = folder + '\\page_' + str(page) + '.zip'
        if not (page + 1) == total_page:
            case_in_page = 20
        else:
            case_in_page = last_page
        name_list = [''] * case_in_page
        id_list = [''] * case_in_page
        date_list = [''] * case_in_page
        print("Download page %s/%s" % (page + 1, total_page))
        for c in range(case_in_page):
            name_list[c] = case_matrix['name'][page + c]
            id_list[c] = case_matrix['doc_id'][page + c]
            date_list[c] = case_matrix['date'][page + c]
        if not os.path.exists(file_name):
            wenshu.downloadDocumentZip(path,
                                       name_list,
                                       id_list,
                                       date_list)
            with open(file_name, "wb") as zipfile:
                zipfile.write(wenshu.doc_content)
            time.sleep(1)
        elif os.path.getsize(file_name) < 30000:
            wenshu.downloadDocumentZip(path,
                                       name_list,
                                       id_list,
                                       date_list)
            with open(file_name, "wb") as zipfile:
                zipfile.write(wenshu.doc_content)
            time.sleep(1)
        else:
            pass
        
    
    
def download_case(search_criteria, court, path):
    folder = path + court
    csv_file = path + court + '.csv'
    print('Read case list from %s' % csv_file)
    # Read csv file and get case list.
    case_matrix = FileUtils.read_csv(csv_file)
    wenshu = Spider.WenShu()
    wenshu.setSearchCriteria(search_criteria)
    
    row_count = len(case_matrix['name'])

    FileUtils.validate_path(folder)
    #for i in range(row_count):
    for i, case_name in enumerate(case_matrix['name']):
        print("%s/%s"%(i, row_count))
        file_name = folder + '\\' + case_name + case_matrix['date'][i] + '.txt'
        if not os.path.exists(file_name):
            wenshu.downloadDocument(path,
                                    case_matrix['name'][i],
                                    case_matrix['doc_id'][i],
                                    case_matrix['date'][i])
            with open(file_name, "wb") as word_doc:
                word_doc.write(wenshu.doc_content)
            time.sleep(1)
        elif os.path.getsize(file_name) < 20000:
            with codecs.open(file_name, "r", "utf-8") as f:
                if not 'DOC' in f.readline():
                    wenshu.downloadDocument(path,
                                            case_matrix['name'][i],
                                            case_matrix['doc_id'][i],
                                            case_matrix['date'][i])
                    with open(file_name, "wb") as word_doc:
                        word_doc.write(wenshu.doc_content)
            time.sleep(1)                        
        else:
            pass
                 
    
def download_invalid(search_criteria, court, path):
    folder = path + court
    csv_file = path + court + '_valid.csv'
    temp_folder = path + court + '_TEMP'
    
    case_matrix = FileUtils.read_csv(csv_file)
    wenshu = Spider.WenShu()
    wenshu.setSearchCriteria(search_criteria)

    #for i in range(row_count):
    for i, case_name in enumerate(case_matrix['name']):
        file_name = folder + '\\' + case_name + case_matrix['date'][i] + '.txt'
        if case_matrix['valid'] == 'N':
            wenshu.downloadDocument(path,
                                    case_matrix['name'][i],
                                    case_matrix['doc_id'][i],
                                    case_matrix['date'][i])
            with open(file_name, "wb") as word_doc:
                word_doc.write(wenshu.doc_content)
            time.sleep(1)
        else:
            pass
    
    file_list = os.listdir(folder)
    FileUtils.validate_path(temp_folder)
    for file in file_list:
        if 'txt' in file:
            move(folder + '\\' + file, temp_folder + '\\' + file[:-4] + '.doc')
   
            
            
def copy_files():
    pass    
    

def validate_csv(path, year):
    csv_file = path + year + '_brief.csv'
    if not os.path.exists(csv_file):
        print("%s not found, please run overview first" % csv_file)
    else:
        br_matrix = FileUtils.read_csv(csv_file)
        c = 0
        for idx, court in enumerate(br_matrix['court']):
            case_matrix = FileUtils.read_csv(path + court + '.csv')
            if not (int(br_matrix['case_number'][idx]) == len(case_matrix['name'])):
                    print("invalid csv file %s" % path + court + '.csv')
                    print('total %s,   excel %s ' % (br_matrix['case_number'][idx], len(case_matrix['name'])))
            #print(idx)    
                
                
def get_total_number(path, year):
    br_matrix = {}
    c = 0
    for key in CourtList.court_list:
        for court in CourtList.court_list[key]:
            c += 1
    
    br_matrix['court'] = [''] * c
    br_matrix['case_number'] = [''] * c
    c = 0
    csv_file = path + year + '_brief.csv'
    if not os.path.exists(csv_file):
        FileUtils.dump2csv(br_matrix, csv_file)
    else:
        br_matrix = FileUtils.read_csv(csv_file)
    wenshu = Spider.WenShu()
    #FileUtils.dump2csv(br_matrix, csv_file)
    for key in CourtList.court_list:
        for court in CourtList.court_list[key]:
            print(c)
            if not br_matrix['court'][c]:
                br_matrix = FileUtils.read_csv(csv_file)
                print('Get total case number of court %s' % court)
                search_criteria = "案件类型:刑事案件,审判程序:一审,法院地域:四川省,裁判年份:" + year +",文书类型:判决书," + "基层法院:" + court
                wenshu.setSearchCriteria(search_criteria)
                wenshu.getTotalItemNumber()
                br_matrix['court'][c] = court
                br_matrix['case_number'][c] = wenshu.total_items
                #print(br_matrix)
                print('Total case number of court %s in %s is %s' % (court, year, br_matrix['case_number'][c]))
                FileUtils.dump2csv(br_matrix, csv_file)
            c += 1

    
    
# Phase 1: Search and get 2nd case list,
#          dump case name list into a csv file.
def download_caselist(search_criteria, csv_file):
    print('Downloading case list of court %s' % csv_file)
    case_matrix = {}

    wenshu = Spider.WenShu()
    wenshu.setSearchCriteria(search_criteria)
    wenshu.getTotalItemNumber()
    print("Total case number is %s" % wenshu.total_items)
    if wenshu.total_items:
        get_case_info(wenshu)
    else:
        print("Failed to get total items.")
        sys.exit(1)
    case_matrix['name'] = wenshu.case['name']
    case_matrix['doc_id'] = wenshu.case['doc_id']
    case_matrix['date'] = wenshu.case['date']
    case_matrix['case_id'] = wenshu.case['case_id']
    case_matrix['procedure'] = wenshu.case['procedure']
    case_matrix['court'] = wenshu.case['court']
    
    print('%s %s' % (wenshu.total_items, len(case_matrix['name'])))
    #row_count = len(case_matrix['name'])
    #pop_count = 0
    #print('row count %s' % row_count)
    #for i in range(row_count):
        #print(i)
    #    if '附带民事' in case_matrix['name2'][i - pop_count]:
            #print('hahaha')
    #        for key in case_matrix:
    #            case_matrix[key].pop(i - pop_count)
    #        pop_count += 1
    FileUtils.dump2csv(case_matrix, csv_file)


#def download_case(search_criteria, csv_file):
    
def phase3(court):
    file_list = os.listdir('.')
    for file in file_list:
        if file[-3:] == 'csv':
            case_matrix = FileUtils.read_csv(file[:-4])
            if len(case_matrix['name']) == 20:
                print(file)
                os.remove(file)

                #print(case_matrix['name'])

def phase5(court_list):
    for court in court_list:
        file_list = os.listdir('Download_' + court)
        for file in file_list:
            copyfile('Download_' + court + '/' + file, 'CasePool/' + file[:-4] + '.doc')
            #if os.path.getsize('Download_' + court + '/' + file) < 100000:
            #    copyfile('Download_' + court + '/' + file, 'CasePool/' + file[:-4] + '.doc')
            #else:
            #    copyfile('Download_' + court + '/' + file, 'CasePool/' + file[:-4] + '.docx')

    
def main():
    desc = "Select a phase to run"
    parser = argparse.ArgumentParser(description=desc)
    parser.add_argument('-d', '--download', action='store_true')
    parser.add_argument('-a', '--analyse', action='store_true')
    parser.add_argument('-t', '--transfer', action='store_true')
    
    parser.add_argument('-v', '--validate', action='store_true')
    parser.add_argument('-y', '--year', action='store')
    parser.add_argument('-c', '--court', action='store')
    parser.add_argument('-r', '--region', action='store')
    parser.add_argument('--count', action='store_true')
    parser.add_argument('--csv', action='store_true')
    parser.add_argument('--overview', action='store_true')
    parser.add_argument('--zip', action='store_true')
    
    
    args = parser.parse_args()
    region = args.region
    year = args.year
    
    path = 'C:\\Users\\lij37\\Code\\Han' + year + '\\'
    #path = 'C:\\Users\\lij37\\Code\\Summer\\2016\\'
    
    
    FileUtils.validate_path(path)
    #search_criteria = "案件类型:刑事案件,审判程序:一审,法院地域:四川省,裁判年份:2016,文书类型:判决书," + "基层法院:" + args.court
    #search_criteria = "案件类型:刑事案件,审判程序:一审,法院地域:四川省,裁判年份:2016,文书类型:判决书,法院层级:中级法院," + "中级法院:" + args.court
    if args.download:
        if args.region:
            for court in CourtList.court_list[region]:
                search_criteria = "案件类型:刑事案件,审判程序:一审,法院地域:四川省,裁判年份:" + year +",文书类型:判决书," + "基层法院:" + court
                csv_file = path + court + '.csv'
                case_folder = path + court
                if args.csv:
                    if not os.path.exists(csv_file):
                        download_caselist(search_criteria, csv_file)
                else:
                    if args.zip:
                        download_case_in_zip(search_criteria, court, path)
                    else:
                        download_case(search_criteria, court, path)    
        elif args.court:
            court = args.court
            search_criteria = "案件类型:刑事案件,审判程序:一审,法院地域:四川省,裁判年份:" + year +",文书类型:判决书," + "基层法院:" + court
            csv_file = path + court + '.csv'
            case_folder = path + court
            if not os.path.exists(csv_file):
                download_caselist(search_criteria, csv_file)
            else:
                if args.zip:
                    download_case_in_zip(search_criteria, court, path)
                else:
                    download_case(search_criteria, court, path)
        else:
            for key in CourtList.court_list:
                for court in CourtList.court_list[key]:
                    search_criteria = "案件类型:刑事案件,审判程序:一审,法院地域:四川省,裁判年份:" + year +",文书类型:判决书," + "基层法院:" + court
                    csv_file = path + court + '.csv'
                    case_folder = path + court
                    if args.csv:
                        if not os.path.exists(csv_file):
                            download_caselist(search_criteria, csv_file)
                    else:
                        if args.zip:
                            download_case_in_zip(search_criteria, court, path)
                        else:
                            download_case(search_criteria, court, path)
    if args.count:
        file_list = os.listdir(path)
        file_count = 0
        for file in file_list:
            if os.path.isdir(path + file):
                count = FileUtils.count_files(path + file)
                print("Total number of docs in %s is %s" % (path + file, count))
                file_count += count
        print("Total number in %s is %s" % (path, file_count))
    
    
    if args.transfer:
        for key in CourtList.court_list:    
            for court in CourtList.court_list[key]:
                print('transfer txt to doc in folder %s' % (path + court))
                FileUtils.transfer2doc(path + court)
    
    if args.analyse:
        for key in CourtList.court_list:    
            for court in CourtList.court_list[key]:
                src = path + 'Download_' + court
                dst = path  + court
                print('copy %s to %s' % (src, dst))
                copytree(src, dst)
    
    if args.validate:
        if args.csv:
            validate_csv(path, year)
    
    if args.overview:
        get_total_number(path, year)            
    
    
        
if __name__ == "__main__":
    main()
